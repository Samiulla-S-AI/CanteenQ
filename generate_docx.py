from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# Add Heading
heading = doc.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run("AGREEMENT FOR CANTEENQ APPLICATION ADOPTION AND OPERATIONS")
run.bold = True
run.font.size = Pt(14)

# Add Date
doc.add_paragraph(f"\nDate: {datetime.date.today().strftime('%B %d, %Y')}")

# Add To section
p = doc.add_paragraph()
p.add_run("To,\n").bold = True
p.add_run("The Manager / Owner,\n[Canteen_name],\n")
p.add_run("Hindusthan Institute of Technology Campus,\nCoimbatore - 641032.\n")

# Add Subject
subject = doc.add_paragraph()
subject.add_run("Subject: ").bold = True
subject.add_run("Agreement for Adoption and Implementation of \"CanteenQ\" Application.")

# Add Salutation
doc.add_paragraph("Dear Sir / Madam,")

# Add Body Paragraph 1
p1 = doc.add_paragraph("We are writing to formally acknowledge your interest in adopting the \"CanteenQ\" application for managing and fulfilling student orders at your canteen. We highly value your willingness to help us promote this application throughout the college campus.")

# Add Body Paragraph 2
p2 = doc.add_paragraph("The CanteenQ application has been fully developed, tested, and will be supported by us, the students of Hindusthan Institute of Technology:")
p2_1 = doc.add_paragraph("1. ")
p2_1.add_run("Vijay M A").bold = True
p2_1.add_run(" (Register No: 720824106124), 2nd Year, B.Tech Information Technology")
p2_2 = doc.add_paragraph("2. ")
p2_2.add_run("Samiulla S").bold = True
p2_2.add_run(" (Register No: 720824108092), 2nd Year, B.Tech Artificial Intelligence and Data Science")

# Add Terms Intro
doc.add_paragraph("\nIn order to ensure smooth operations and mutual understanding, we hereby agree to the following conditions:")

# Add Term 1
t1 = doc.add_paragraph()
t1.add_run("1. Technical & Physical Responsibility: ").bold = True
t1.add_run("As the sole developers of the CanteenQ application, we (Vijay M A and Samiulla S) take full and complete responsibility for resolving any technical or physical issues that may occur during the utilization of the app within the campus.")

# Add Term 2
t2 = doc.add_paragraph()
t2.add_run("2. Daily Settlement of Funds: ").bold = True
t2.add_run("All the money collected from students for their completed food orders via the CanteenQ application will be consolidated and handed over to you on the same day. Specifically, the total daily amount will be provided to [Canteen_name] every evening after college hours, post 4:45 PM.")

# Add concluding paragraph
doc.add_paragraph("\nWe are excited to partner with you and look forward to improving the canteen experience for all students and staff.")

doc.add_paragraph("Yours Sincerely,\n")

# Add Signatures Section
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True

cell1 = table.cell(0, 0)
p_cell1 = cell1.paragraphs[0]
p_cell1.add_run("______________________________\n").bold = True
p_cell1.add_run("Vijay M A\n")
p_cell1.add_run("2nd Year, B.Tech IT\n")
p_cell1.add_run("Hindusthan Institute of Technology")

cell2 = table.cell(0, 1)
p_cell2 = cell2.paragraphs[0]
p_cell2.add_run("______________________________\n").bold = True
p_cell2.add_run("Samiulla S\n")
p_cell2.add_run("2nd Year, B.Tech AI & DS\n")
p_cell2.add_run("Hindusthan Institute of Technology")

# Add Acknowledgment section
ack_heading = doc.add_paragraph("\n\nACKNOWLEDGMENT & ACCEPTANCE")
ack_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
ack_heading.runs[0].bold = True
ack_heading.runs[0].font.size = Pt(12)

doc.add_paragraph("I, representing [Canteen_name], hereby accept the terms outlined above and confirm our participation in adopting the CanteenQ application. We also agree to support its promotion across the college.")

doc.add_paragraph("\n\n______________________________\nSignature with Canteen Seal / Stamp\nName:\nDate:")

doc.save("CanteenQ_Agreement_v2.docx")
