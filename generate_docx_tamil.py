from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# Add Heading
heading = doc.add_paragraph()
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = heading.add_run("CanteenQ செயலி அமலாக்கம் மற்றும் செயல்பாட்டு ஒப்பந்தம்")
run.bold = True
run.font.size = Pt(14)

# Add Date
doc.add_paragraph(f"\nதேதி: {datetime.date.today().strftime('%B %d, %Y')}")

# Add To section
p = doc.add_paragraph()
p.add_run("பெறுநர்,\n").bold = True
p.add_run("மேலாளர் / உரிமையாளர்,\n[Canteen_name],\n")
p.add_run("ஹிந்துஸ்தான் இன்ஸ்டிடியூட் ஆப் டெக்னாலஜி வளாகம்,\nகோயம்புத்தூர் - 641032.\n")

# Add Subject
subject = doc.add_paragraph()
subject.add_run("பொருள்: ").bold = True
subject.add_run("\"CanteenQ\" செயலியை ஏற்றுக்கொள்வது மற்றும் செயல்படுத்துவது தொடர்பான ஒப்பந்தம்.")

# Add Salutation
doc.add_paragraph("மதிப்பிற்குரிய ஐயா / அம்மா,")

# Add Body Paragraph 1
p1 = doc.add_paragraph("உங்கள் உணவகத்தில் மாணவர்களின் ஆர்டர்களை நிர்வகிக்கவும் பூர்த்தி செய்யவும் \"CanteenQ\" செயலியை பயன்படுத்த நீங்கள் காட்டிய ஆர்வத்திற்கு எங்கள் முறையான நன்றியைத் தெரிவித்துக் கொள்கிறோம். கல்லூரி வளாகம் முழுவதும் இந்தச் செயலியைப் பயன்படுத்த நீங்கள் அளிக்கும் ஆதரவை நாங்கள் பெரிதும் மதிக்கிறோம்.")

# Add Body Paragraph 2
p2 = doc.add_paragraph("இந்த CanteenQ செயலி முழுமையாக உருவாக்கப்பட்டு, சோதிக்கப்பட்டுள்ளது, மேலும் ஹிந்துஸ்தான் இன்ஸ்டிடியூட் ஆப் டெக்னாலஜி மாணவர்களான எங்களால் நிர்வகிக்கப்படும்:")
p2_1 = doc.add_paragraph("1. ")
p2_1.add_run("விஜய் M A").bold = True
p2_1.add_run(" (பதிவு எண்: 720824106124), 2-ம் ஆண்டு, B.Tech தகவல் தொழில்நுட்பம்")
p2_2 = doc.add_paragraph("2. ")
p2_2.add_run("சமியுல்லா S").bold = True
p2_2.add_run(" (பதிவு எண்: 720824108092), 2-ம் ஆண்டு, B.Tech செயற்கை நுண்ணறிவு மற்றும் தரவு அறிவியல்")

# Add Terms Intro
doc.add_paragraph("\nசுமூகமான செயல்பாடுகள் மற்றும் பரஸ்பர புரிதலை உறுதி செய்ய, பின்வரும் நிபந்தனைகளுக்கு நாங்கள் ஒப்புக்கொள்கிறோம்:")

# Add Term 1
t1 = doc.add_paragraph()
t1.add_run("1. தொழில்நுட்பம் மற்றும் நேரடிப் பொறுப்பு: ").bold = True
t1.add_run("CanteenQ செயலியின் முழுமையான உருவாக்குநர்களாகிய நாங்கள் (விஜய் M A மற்றும் சமியுல்லா S), வளாகத்திற்குள் இந்தச் செயலியைப் பயன்படுத்தும்போது ஏற்படும் எந்தவொரு தொழில்நுட்ப அல்லது பிற பிரச்சனைகளையும் தீர்க்க முழுப் பொறுப்பேற்கிறோம்.")

# Add Term 2
t2 = doc.add_paragraph()
t2.add_run("2. தினசரி நிதி தீர்வு: ").bold = True
t2.add_run("CanteenQ செயலி மூலம் மாணவர்களின் உணவுகளுக்காகப் பெறப்படும் அனைத்து தொகையும் ஒவ்வொரு நாளும் மாலை மொத்தமாக உங்களிடம் வழங்கப்படும். குறிப்பாக, ஒவ்வொரு நாளும் கல்லூரி முடிந்த பிறகு, மாலை 4:45 மணிக்கு மேல் [Canteen_name]-க்கு அன்றைய முழுத் தொகையும் வழங்கப்படும்.")

# Add concluding paragraph
doc.add_paragraph("\nஉங்களுடன் இணைந்து செயல்படுவதில் நாங்கள் மகிழ்ச்சியடைகிறோம், மேலும் மாணவர்கள் மற்றும் பணியாளர்கள் அனைவருக்கும் கேண்டீன் அனுபவத்தை மேம்படுத்த ஆவலுடன் காத்திருக்கிறோம்.")

doc.add_paragraph("இப்படிக்கு,\n")

# Add Signatures Section
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True

cell1 = table.cell(0, 0)
p_cell1 = cell1.paragraphs[0]
p_cell1.add_run("______________________________\n").bold = True
p_cell1.add_run("விஜய் M A\n")
p_cell1.add_run("2-ம் ஆண்டு, B.Tech IT\n")
p_cell1.add_run("ஹிந்துஸ்தான் இன்ஸ்டிடியூட் ஆப் டெக்னாலஜி")

cell2 = table.cell(0, 1)
p_cell2 = cell2.paragraphs[0]
p_cell2.add_run("______________________________\n").bold = True
p_cell2.add_run("சமியுல்லா S\n")
p_cell2.add_run("2-ம் ஆண்டு, B.Tech AI & DS\n")
p_cell2.add_run("ஹிந்துஸ்தான் இன்ஸ்டிடியூட் ஆப் டெக்னாலஜி")

# Add Acknowledgment section
ack_heading = doc.add_paragraph("\n\nஒப்புதல் மற்றும் ஏற்பு")
ack_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
ack_heading.runs[0].bold = True
ack_heading.runs[0].font.size = Pt(12)

doc.add_paragraph("[Canteen_name] சார்பில் நான், மேற்கண்ட நிபந்தனைகளை ஏற்றுக்கொண்டு CanteenQ செயலியைப் பயன்படுத்துவதில் எங்கள் முழுமையான பங்களிப்பை மீண்டும் உறுதிப்படுத்துகிறேன். கல்லூரி முழுவதும் இந்தச் செயலியை விரிவுபடுத்த ஆதரவளிக்கச் சம்மதிக்கிறோம்.")

doc.add_paragraph("\n\n______________________________\nகேண்டீன் முத்திரையுடன் கையொப்பம்\nபெயர்:\nதேதி:")

doc.save("CanteenQ_Agreement_Tamil.docx")
